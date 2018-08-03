#!/usr/bin/python3
import json
import argparse
import docx
import docxtra
from read_json import *

def output_paper(paper, para, authors, venues):
    print(paper['id'])
    # Authors
    for a in paper['authorIds']:
        if a in authors:
            a = authors[a]
        para.add_run(a + ", ")
    # Title
    para.add_run(paper['title'] + ". ").bold = True
    # Venue
    vstr = paper['venueId']
    if vstr in venues:
        vstr = venues[vstr]
    para.add_run(vstr + ", ").italic = True
    # year
    para.add_run(str(paper['year']))
    # link
    para.add_run(" (")
    link = paper['pdfLink']
    if link.startswith("files"):
        link = "http://sameersingh.org/" + link
    docxtra.add_hyperlink(para, link, 'link', 'FF8822', False)
    para.add_run(")")

def run(idir, ofile):
    pfile = idir + "/papers.json"
    afile = idir + "/authors.json"
    vfile = idir + "/venues.json"
    authors = read_authors(afile)
    venues = read_venues(vfile)
    papers = read_papers(pfile)
    document = docx.Document()
    document.add_heading('Publications', 0)
    # Book Chapters
    document.add_heading('Book Chapter', 1)
    for p in papers:
        if p['pubTypeSlot'] == 'Chapter':
            para = document.add_paragraph('', style='ListBullet')
            output_paper(p, para, authors, venues)
    # Patents
    document.add_heading('Patents', 1)
    for p in papers:
        if p['pubTypeSlot'] == 'Patent':
            para = document.add_paragraph('', style='ListBullet')
            output_paper(p, para, authors, venues)
    # Journals
    document.add_heading('Journal', 1)
    for p in papers:
        if p['pubTypeSlot'] == 'Journal':
            para = document.add_paragraph('', style='ListBullet')
            output_paper(p, para, authors, venues)
    # Conference
    document.add_heading('Peer-Reviewed Conference', 1)
    for p in papers:
        if p['pubTypeSlot'] == 'Conference' and 'pdfLink' in p:
            para = document.add_paragraph('', style='ListBullet')
            output_paper(p, para, authors, venues)
    #print(json.dumps(data))
    document.save(ofile)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-i", "--input", help="directory containing the json files for authors/papers", required=True)
    parser.add_argument("-o", "--output", help="docx file to write the papers in", required=True)
    args = parser.parse_args()
    print("input:  ", args.input)
    print("output: ", args.output)
    run(args.input, args.output)
